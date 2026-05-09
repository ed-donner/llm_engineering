"""
Week 1 Community Contribution
Hierarchical Document Summarizer
Cloud (gpt-4o-mini) with Ollama fallback (llama3.2:3b)
"""

import os
import time
import logging
import argparse
from openai import OpenAI
from docx import Document
from PyPDF2 import PdfReader
import docx2txt


# =====================================================
# LOGGING
# =====================================================

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s"
)


# =====================================================
# FILE READING
# =====================================================

def read_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    elif ext == ".pdf":
        text = ""
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text

    elif ext == ".docx":
        return docx2txt.process(file_path)

    else:
        raise ValueError("Unsupported file type")


# =====================================================
# SMART CHUNKING
# =====================================================

def split_text_safely(text, chunk_size=5000):
    chunks = []
    while len(text) > chunk_size:
        split_at = text.rfind(".", 0, chunk_size)
        if split_at == -1:
            split_at = chunk_size
        chunks.append(text[:split_at])
        text = text[split_at:]
    chunks.append(text)
    return chunks


# =====================================================
# CLIENTS
# =====================================================

def get_cloud_client():
    key = os.getenv("OPENAI_API_KEY")
    if not key:
        logging.warning("OPENAI_API_KEY not found. Using local model.")
        return None
    return OpenAI(api_key=key, timeout=300)


def get_local_client():
    return OpenAI(
        base_url="http://localhost:11434/v1",
        api_key="ollama",
        timeout=600
    )


# =====================================================
# SAFE EXTRACTION
# =====================================================

def extract_content(response):
    if not response:
        return ""

    if hasattr(response, "output_text") and response.output_text:
        return response.output_text

    try:
        return response.choices[0].message.content or ""
    except Exception:
        return ""


# =====================================================
# SAFE COMPLETION
# =====================================================

def safe_completion(messages, max_tokens, cloud_client, local_client):

    model_cloud = "gpt-4o-mini"
    model_local = "llama3.2:3b"

    if cloud_client:
        try:
            logging.info("Using CLOUD model...")
            return cloud_client.chat.completions.create(
                model=model_cloud,
                messages=messages,
                max_completion_tokens=max_tokens
            )
        except Exception as e:
            logging.warning(f"Cloud failed: {e}")

    if local_client:
        try:
            logging.info("Using LOCAL model...")
            return local_client.chat.completions.create(
                model=model_local,
                messages=messages,
                max_tokens=max_tokens,
                temperature=0.2
            )
        except Exception as e:
            logging.error(f"Local failed: {e}")

    raise RuntimeError("Both cloud and local completions failed.")


# =====================================================
# SUMMARIZATION
# =====================================================

SYSTEM_PROMPT = "You are an expert technical document summarizer."

def summarize_text(text, batch_size=4):

    cloud_client = get_cloud_client()
    local_client = get_local_client()

    chunks = split_text_safely(text)
    logging.info(f"Total chunks: {len(chunks)}")

    chunk_summaries = []

    for i, chunk in enumerate(chunks):
        logging.info(f"Summarizing chunk {i+1}/{len(chunks)}")

        messages = [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": f"""
Summarize this section and extract:
- Key themes
- Strategic objectives
- Risks
- Recommendations
- Do not translate the language.

TEXT:
{chunk}
"""}
        ]

        response = safe_completion(messages, 600, cloud_client, local_client)
        content = extract_content(response)

        if content.strip():
            chunk_summaries.append(content)

    if not chunk_summaries:
        return "All chunk summaries failed."

    batch_summaries = []

    for i in range(0, len(chunk_summaries), batch_size):
        batch = "\n\n".join(chunk_summaries[i:i+batch_size])

        messages = [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": f"""
Create a concise executive summary.

CONTENT:
{batch}
"""}
        ]

        response = safe_completion(messages, 800, cloud_client, local_client)
        content = extract_content(response)

        if content.strip():
            batch_summaries.append(content)

    final_input = "\n\n".join(batch_summaries)

    final_messages = [
        {"role": "system", "content": "You create structured leadership briefings."},
        {"role": "user", "content": f"""
Create a final executive summary.

CONTENT:
{final_input}
"""}
    ]

    final_response = safe_completion(final_messages, 1000, cloud_client, local_client)
    return extract_content(final_response)


# =====================================================
# SAVE DOCX
# =====================================================

def save_summary(summary_text, output_path):
    doc = Document()
    doc.add_heading("Executive Summary", level=1)

    for line in summary_text.split("\n"):
        doc.add_paragraph(line)

    doc.save(output_path)


# =====================================================
# MAIN (CLI)
# =====================================================

if __name__ == "__main__":

    parser = argparse.ArgumentParser(description="Hierarchical Document Summarizer")
    parser.add_argument("input_file", help="Path to input document (.txt, .pdf, .docx)")
    parser.add_argument("--output_dir", default=".", help="Output directory")

    args = parser.parse_args()

    input_file = args.input_file
    output_dir = args.output_dir

    os.makedirs(output_dir, exist_ok=True)

    base_name = os.path.splitext(os.path.basename(input_file))[0]
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    output_file = os.path.join(output_dir, f"{base_name}_summary_{timestamp}.docx")

    logging.info("Reading document...")
    text = read_file(input_file)

    logging.info("Generating summary...")
    summary = summarize_text(text)

    logging.info("Saving summary...")
    save_summary(summary, output_file)

    logging.info(f"Done. Saved at: {output_file}")