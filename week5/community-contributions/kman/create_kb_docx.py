"""Create Knowledge_Base_Guide.docx in knowledge_base folder for the Week 5 RAG exercise."""
import os
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Install python-docx: pip install python-docx")
    raise

KB_DIR = "knowledge_base"
os.makedirs(KB_DIR, exist_ok=True)
path = os.path.join(KB_DIR, "Knowledge_Base_Guide.docx")

doc = Document()
doc.add_heading("Personal Knowledge Base — Guide", 0)
doc.add_paragraph(
    "This document is part of the Week 5 exercise: a personal knowledge worker that uses RAG (Retrieval Augmented Generation) to answer questions from your own files."
)
doc.add_heading("What is in this knowledge base?", level=1)
doc.add_paragraph(
    "The knowledge base can contain text files (.txt), Markdown (.md), PDFs (.pdf), and Word documents (.docx). "
    "All documents are chunked, embedded with sentence-transformers, and stored in Chroma. When you ask a question, "
    "the system retrieves the most relevant chunks and uses an LLM (via OpenRouter) to generate an answer grounded in your content."
)
doc.add_heading("How to use the RAG chat", level=1)
doc.add_paragraph(
    "Run the notebook cells in order: (1) Imports and setup, (2) Load documents, (3) Build Chroma vector store, (4) Define RAG chain, (5) Launch Gradio chat. "
    "Then type questions in natural language. Examples: \"What is in my knowledge base?\", \"How do I add Word documents?\", \"What is RAG?\""
)
doc.add_heading("Adding more documents", level=1)
doc.add_paragraph(
    "Place any .txt, .md, .pdf, or .docx file in the knowledge_base folder. Re-run the document loading cell and the Chroma cell to re-index. "
    "The vector store will then include the new content for retrieval."
)
doc.add_heading("Technical stack", level=1)
doc.add_paragraph(
    "Embeddings: HuggingFace sentence-transformers (all-MiniLM-L6-v2), no API key required. "
    "Vector store: Chroma (persisted locally). LLM: OpenRouter (e.g. Gemini) using OPENROUTER_API_KEY. "
    "UI: Gradio ChatInterface."
)
doc.save(path)
print(f"Created: {os.path.abspath(path)}")
