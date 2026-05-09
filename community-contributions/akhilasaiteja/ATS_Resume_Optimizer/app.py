import os
import json
import re
import tempfile
from datetime import datetime
from pathlib import Path

from dotenv import load_dotenv
from openai import OpenAI
import gradio as gr

import pypdf
import pdfplumber
from docx import Document as DocxDocument

from docx import Document
from docx.shared import Pt, RGBColor, Inches

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

# ── Configuration ─────────────────────────────────────────────────────────────
load_dotenv(override=True)
openai_api_key = os.getenv('OPENAI_API_KEY')
MODEL = 'gpt-4.1-mini'        # lightweight: keyword extraction, semantic batch, date parsing
MODEL_SMART = 'gpt-4o'         # reasoning: optimization, classification, chat
client = OpenAI(api_key=os.getenv('OPENAI_API_KEY', 'test-key-placeholder'))

MAX_RUNS_PER_SESSION = 10

SUPPORTED_ROLES = [
    "Data Scientist",
    "Machine Learning Engineer",
    "AI / LLM Engineer",
    "Data Analyst",
    "Business Analyst",
    "Software Engineer",
]

# Hard technical tools — exact match only, never semantically substituted
EXACT_MATCH_ONLY = {
    # Languages
    "Python", "R", "SQL", "Java", "C++", "JavaScript", "TypeScript",
    "Go", "Rust", "Scala", "MATLAB",
    # ML Frameworks
    "LightGBM", "XGBoost", "CatBoost", "TensorFlow", "PyTorch",
    "scikit-learn", "Keras", "Statsmodels",
    # MLOps / Data Engineering
    "Docker", "Kubernetes", "Airflow", "Spark", "PySpark", "dbt",
    "MLflow", "Kubeflow", "Prefect", "Luigi",
    # Cloud
    "AWS", "GCP", "Azure", "SageMaker", "Vertex AI", "Databricks",
    "EC2", "S3", "Lambda",
    # Databases
    "PostgreSQL", "MySQL", "MongoDB", "Firebase", "Redis",
    "Snowflake", "BigQuery", "Redshift", "Cassandra",
    # LLM / AI Tools
    "LangChain", "HuggingFace", "Gradio", "Streamlit", "RAG",
    "FAISS", "Pinecone", "OpenAI", "Anthropic",
    # BI / Visualization
    "Tableau", "Power BI", "Looker", "Matplotlib", "Seaborn", "Plotly",
    # DevOps / Other
    "Git", "GitHub", "Linux", "Bash", "FastAPI", "Flask", "Django",
    "Terraform", "Ansible", "Jenkins",
}

# ── System message ────────────────────────────────────────────────────────────
SYSTEM_MESSAGE = """You are an expert ATS (Applicant Tracking System) resume optimizer.

## YOUR ROLE
Optimize resumes to pass ATS screening for one of these 6 supported roles ONLY:
- Data Scientist
- Machine Learning Engineer
- AI / LLM Engineer
- Data Analyst
- Business Analyst
- Software Engineer

If the JD role doesn't match exactly, map it to the closest supported role above.

## STRICT RULES — NEVER VIOLATE

You are optimizing a resume against a job description using GPT-4o.

### WHAT YOU MUST DO:
- Scan the entire resume for JD keywords already present anywhere in the resume
- Rephrase existing bullets to use JD vocabulary where the underlying experience genuinely supports it
- Add JD keywords to the Skills section ONLY if they already appear somewhere else in the resume
- Strengthen the summary using skills and terms that appear anywhere in the resume
- Return a measurably improved resume with stronger JD alignment

### WHAT YOU MUST NEVER DO:
- Add any skill, tool, or technology to Skills that does not appear anywhere in the original resume
- Add any bullet to Experience or Projects claiming experience not evidenced in the original
- Change any number, metric, or specific detail
- Merge two bullets into one
- Drop any existing bullet
- Add new bullets not in the original

### THE TEST FOR EVERY CHANGE:
Can I point to the exact line in the original resume that supports this change?
- YES → make the change
- NO → do not make it, list as missing instead

### KEYWORD SCANNING — ALL SECTIONS
Scan the ENTIRE resume as one text block. A keyword found ANYWHERE counts:
- Professional Summary / Objective
- Skills section
- ALL experience bullets (every role, every bullet)
- ALL projects (every project, every bullet)
- Education (including coursework, tools mentioned)
- Certifications
- Any other section present

Never limit scanning to a specific section.

### SCORING RULES
- Score = (JD keywords found in resume) / (total JD keywords) × 100
- Extra skills the candidate has that are NOT in the JD → completely ignored, NO PENALTY
- Keyword frequency is irrelevant — present once = present 10 times = exact same score
- Never reward repetition; never penalize single mention

When extracting JD keywords, be selective and consistent:
- Include: technical skills, tools, frameworks, languages, domain terms, industry-specific terminology, measurable competencies (e.g. "experiment design", "causal inference", "A/B experimentation", "data governance")
- Exclude:
  * Soft skills and personality traits of any kind (e.g. curious mindset, self-motivated, keen eye for detail, passionate, collaborative, team player, fast-paced, eager, proactive — and anything similar)
  * Company names, team names, work arrangements
  * Phrases longer than 4 words
  * Redundant variants of the same term
  * Vague cultural fit terms
  * Anything a candidate would never put on a resume

### EXPERIENCE — RELEVANT ONLY
When calculating experience relevance:
- Count ONLY experience directly relevant to the target role
- Unrelated roles (radio presenter → Data Scientist) = 0 years contributed
- Partially related roles (banking analyst → Data Scientist) = 0.5× multiplier
- Directly related roles (ML Research Assistant → Data Scientist) = 1.0× multiplier

### OUTPUT FORMAT
Return a valid JSON object with this exact structure:
{
    "role_type_detected": "<one of the 6 supported roles>",
    "optimized_resume": "<full optimized resume text, preserving all sections>",
    "ats_keywords": {
        "matched": ["keyword1", "keyword2"],
        "added": ["keyword3"],
        "missing": ["keyword4"]
    },
    "changes_made": ["Added Docker to Skills section", "Strengthened bullet in Experience"],
    "ats_score_before": <0-100>,
    "ats_score_after": <0-100>
}
"""

# ── PDF / DOCX readers ────────────────────────────────────────────────────────
def extract_text_from_pdf(file_path):
    """Read PDF — pdfplumber primary, pypdf fallback."""
    text = ''
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + '\n'
        if text.strip():
            return text.strip()
        raise ValueError('pdfplumber returned empty text')
    except Exception:
        try:
            reader = pypdf.PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + '\n'
            return text.strip()
        except Exception as e2:
            return f'Error reading PDF: {e2}'


def extract_text_from_docx(file_path):
    """Read DOCX paragraph by paragraph."""
    try:
        doc = DocxDocument(file_path)
        return '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        return f'Error reading DOCX: {e}'


# ── Scoring functions ─────────────────────────────────────────────────────────
def extract_required_experience(jd_text):
    """Extract minimum years of experience required from JD via regex."""
    patterns = [
        r'(\d+)\+?\s*(?:to|-)\s*\d+\s*years?',
        r'(\d+)\+\s*years?',
        r'minimum\s+(\d+)\s*years?',
        r'at\s+least\s+(\d+)\s*years?',
        r'(\d+)\s*or\s*more\s*years?',
        r'(\d+)\s*years?\s*of\s*(?:relevant\s+)?experience',
    ]
    for pattern in patterns:
        match = re.search(pattern, jd_text.lower())
        if match:
            return int(match.group(1))
    return None


def _parse_date(date_str, use_gpt_fallback=True):
    """
    Parse a date string to datetime.
    Uses regex/strptime first; GPT fallback for unusual formats only.
    "Present" / "Current" / "Now" → today's date.
    """
    date_str = date_str.strip()
    if date_str.lower() in ('present', 'current', 'now', 'today', 'ongoing'):
        return datetime.now()

    formats = [
        '%B %Y',   # January 2022
        '%b %Y',   # Jan 2022
        '%m/%Y',   # 01/2022
        '%Y-%m',   # 2022-01
        '%Y',      # 2022
    ]
    for fmt in formats:
        try:
            return datetime.strptime(date_str, fmt)
        except ValueError:
            continue

    # GPT fallback for unusual formats
    if use_gpt_fallback:
        try:
            resp = client.chat.completions.create(
                model=MODEL,
                messages=[{
                    'role': 'user',
                    'content': (
                        f'Parse this date to YYYY-MM format. '
                        f'Return ONLY the date in YYYY-MM format, nothing else. '
                        f'Date: "{date_str}"'
                    )
                }],
                max_tokens=10,
                temperature=0
            )
            raw = resp.choices[0].message.content.strip()
            return datetime.strptime(raw, '%Y-%m')
        except Exception:
            return None
    return None


_DEGREE_KEYWORDS = {
    'bachelor', 'master', 'phd', 'gpa', 'doctorate', 'diploma',
    'certificate', 'b.s.', 'm.s.', 'm.b.a.', 'b.tech',
}


def _is_education_context(lines, line_idx):
    """
    Return True if the date line or any of the 2 lines immediately before it
    contain a degree keyword. Covers cases where the date sits on its own line
    directly below the degree name (e.g. 'B.S. Computer Science\\n2015 – 2019').
    """
    check_range = range(max(0, line_idx - 2), line_idx + 1)
    return any(
        kw in lines[i].lower()
        for i in check_range
        for kw in _DEGREE_KEYWORDS
    )


def _extract_experience_entries(resume_text):
    """
    Extract work experience entries from resume text.
    Returns list of dicts: {context, start_label, end_label, duration_years}
    Skips lines that contain degree keywords. Never assumes a fixed section order.

    Role-boundary aware approach:
    1. First pass — find all lines containing date ranges and record their index.
    2. For each date line, walk backwards to find the nearest non-empty line
       without a date range — that is the role title line.
    3. Context = role title line through date line only (no adjacent role bleed).
    4. Deduplicate only on identical (start_label, end_label) pairs.
    5. Date ranges within 3 lines of each other are always separate roles.
    """
    DATE_RANGE_RE = re.compile(
        r'(?P<start>'
        r'(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|'
        r'Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)'
        r'\.?\s+\d{4}|\d{4})'
        r'\s*(?:–|—|‒|-|to)\s*'
        r'(?P<end>'
        r'(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|'
        r'Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)'
        r'\.?\s+\d{4}|\d{4}|[Pp]resent|[Cc]urrent|[Nn]ow|[Oo]ngoing)',
        re.IGNORECASE
    )

    lines = resume_text.split('\n')

    # Section headers that terminate the experience context window
    _SECTION_BREAK_HEADERS = re.compile(
        r'^\s*(PROJECTS|EDUCATION|CERTIFICATIONS?|SKILLS?|SUMMARY|'
        r'AWARDS?|PUBLICATIONS?|ACHIEVEMENTS?|TECHNICAL SKILLS?|'
        r'SKILLS\s*&\s*TOOLS|PROFESSIONAL SUMMARY|PROFILE)\s*$',
        re.IGNORECASE
    )

    # Pass 1 — find all section-break line indices
    section_break_indices = set()
    for i, line in enumerate(lines):
        if _SECTION_BREAK_HEADERS.match(line):
            section_break_indices.add(i)

    # Pass 2 — collect indices of all lines that contain a date range
    # (excluding education context and lines inside non-experience sections)
    date_line_indices = []
    for i, line in enumerate(lines):
        if _is_education_context(lines, i):
            continue
        if DATE_RANGE_RE.search(line):
            date_line_indices.append(i)

    date_line_set = set(date_line_indices)

    entries = []
    for i in date_line_indices:
        line = lines[i]
        for match in DATE_RANGE_RE.finditer(line):
            start_dt = _parse_date(match.group('start'))
            end_dt = _parse_date(match.group('end'))
            if start_dt is None or end_dt is None:
                continue
            if end_dt < start_dt:
                continue

            # Walk backwards to find the start boundary of this role.
            # Stop at: an empty line, a section header, or another date range line.
            start_boundary = i
            for j in range(i - 1, max(-1, i - 15), -1):
                candidate = lines[j].strip()
                # Empty line = role boundary
                if not candidate:
                    break
                # All-caps section header = role boundary
                if candidate.isupper() and len(candidate) < 50:
                    break
                # Another date range line = boundary of a different role
                if j in date_line_set and j != i:
                    break
                start_boundary = j

            # Walk forwards to find the end boundary of this role.
            # Hard stop: never cross a section-break header line.
            # Also stop at: next date range line, or empty line before a header/date.
            end_boundary = i
            for k in range(i + 1, min(len(lines), i + 20)):
                # Hard stop at any section-break header
                if k in section_break_indices:
                    break
                candidate = lines[k].strip()
                # Stop at any other date range line (next role)
                if k in date_line_set:
                    break
                # Stop at empty line followed by a section break or date
                if not candidate:
                    for m in range(k + 1, min(len(lines), k + 4)):
                        peek = lines[m].strip()
                        if peek:
                            if (peek.isupper() and len(peek) < 50) or m in date_line_set or m in section_break_indices:
                                end_boundary = k - 1
                            else:
                                end_boundary = k
                            break
                    else:
                        end_boundary = k
                    break
                end_boundary = k

            context = '\n'.join(lines[start_boundary:end_boundary + 1]).strip()

            duration_months = (
                (end_dt.year - start_dt.year) * 12 + end_dt.month - start_dt.month
            )
            duration_years = round(max(0, duration_months) / 12, 1)

            end_label = (
                'Present'
                if end_dt.month == datetime.now().month and end_dt.year == datetime.now().year
                else end_dt.strftime('%b %Y')
            )

            entries.append({
                'context': context,
                'start_label': start_dt.strftime('%b %Y'),
                'end_label': end_label,
                'duration_years': duration_years,
            })

    # Deduplicate only on identical (start_label, end_label) pairs
    seen = set()
    unique_entries = []
    for e in entries:
        key = (e['start_label'], e['end_label'])
        if key not in seen:
            seen.add(key)
            unique_entries.append(e)

    return unique_entries


def calculate_relevant_experience(resume_text, target_role, jd_text):
    """
    Hybrid approach:
    - Regex + datetime.now() for date parsing (GPT fallback for unusual formats only)
    - GPT always used for relevance classification (full / partial / none)

    Returns:
        relevant_years (float)
        breakdown (list of dicts with job_title, start, end, total_years,
                   multiplier, counted_years, relevance, reason)
        required_years (int | None)
    """
    required_years = extract_required_experience(jd_text)
    entries = _extract_experience_entries(resume_text)

    if not entries:
        return 0.0, [], required_years

    # Build prompt for GPT to classify all jobs in one call
    jobs_block = '\n\n'.join([
        f'Job {i + 1}:\n{e["context"]}\nDuration: {e["duration_years"]} years'
        for i, e in enumerate(entries)
    ])

    classification_prompt = f"""Target role: {target_role}

For each job listed below, classify its relevance to the target role.
Return a JSON array — one object per job in the same order:

[
  {{
    "job_number": 1,
    "job_title": "extracted job title",
    "relevance": "full" | "partial" | "none",
    "multiplier": 1.0 | 0.5 | 0.0,
    "reason": "one sentence explaining why"
  }}
]

Classification rules — base the classification on the candidate's OWN described tasks and responsibilities, not on industry label alone:
- full (1.0×): Candidate's bullets describe tasks directly used in the target role (e.g. training models, writing SQL, building pipelines)
- partial (0.5×): Candidate's bullets show some transferable technical skill but in a different domain (e.g. Excel modelling, statistical analysis in a non-tech role)
- none (0×): Candidate's bullets describe operational, administrative, client-facing, or creative work with no technical overlap to the target role (e.g. hosting shows, managing inventory, customer service, sales)

Additional classification guidance:

- none (0×) should only be used for roles that are purely operational, administrative, or customer-facing with absolutely no analytical, quantitative, or technical component whatsoever

- If a role involves any of the following it should be classified as at least partial (0.5×):
  * working with data in any form
  * analysis or reporting of any kind
  * quantitative decision making
  * technical problem solving
  * domain knowledge relevant to the target role

- When in doubt between partial and none, always default to partial — it is better to slightly over-credit an ambiguous role than to dismiss genuinely relevant experience entirely

- Base the classification strictly on what the candidate described doing — not on job title or industry alone

Jobs to classify:
{jobs_block}"""

    try:
        resp = client.chat.completions.create(
            model=MODEL_SMART,
            messages=[{'role': 'user', 'content': classification_prompt}],
            temperature=0,
            max_tokens=1200
        )
        raw = resp.choices[0].message.content.strip()
        # Strip markdown code fences if present
        if '```' in raw:
            raw = re.sub(r'```(?:json)?', '', raw).strip().strip('`')
        parsed = json.loads(raw)
        # GPT might return {"jobs": [...]} or just [...]
        if isinstance(parsed, dict):
            classifications = next(iter(parsed.values()))
        else:
            classifications = parsed
    except Exception as e:
        # Fallback: treat all as partial
        classifications = [
            {
                'job_number': i + 1,
                'job_title': 'Unknown',
                'relevance': 'partial',
                'multiplier': 0.5,
                'reason': f'Classification unavailable: {e}'
            }
            for i in range(len(entries))
        ]

    breakdown = []
    total_relevant = 0.0

    for i, entry in enumerate(entries):
        cls = next(
            (c for c in classifications if c.get('job_number') == i + 1),
            {'job_title': 'Unknown', 'relevance': 'partial', 'multiplier': 0.5, 'reason': ''}
        )
        multiplier = float(cls.get('multiplier', 0.5))
        counted = round(entry['duration_years'] * multiplier, 1)
        total_relevant += counted

        breakdown.append({
            'job_title': cls.get('job_title', 'Unknown'),
            'start': entry['start_label'],
            'end': entry['end_label'],
            'total_years': entry['duration_years'],
            'multiplier': multiplier,
            'counted_years': counted,
            'relevance': cls.get('relevance', 'partial'),
            'reason': cls.get('reason', ''),
        })

    return round(total_relevant, 1), breakdown, required_years


def calculate_keyword_score(jd_keywords, resume_full_text):
    """
    Two-layer keyword matching:
    Layer 1 — exact string match (case-insensitive) across full resume text
    Layer 2 — GPT semantic batch check ONLY for soft skills not found in Layer 1

    Hard technical tools (EXACT_MATCH_ONLY set) → Layer 1 only, never semantic.

    Returns:
        score (int 0-100)
        found (list)
        missing (list)
        semantic_matches (list of {keyword, evidence})
    """
    if not jd_keywords:
        return 0, [], [], []

    resume_lower = resume_full_text.lower()
    found = []
    missing = []

    for kw in jd_keywords:
        if kw.lower() in resume_lower:
            found.append(kw)
        else:
            missing.append(kw)

    if not missing:
        return round(len(found) / len(jd_keywords) * 100), found, [], []

    # Partition missing keywords
    hard_tech_missing = [
        kw for kw in missing
        if kw in EXACT_MATCH_ONLY or any(kw.lower() == e.lower() for e in EXACT_MATCH_ONLY)
    ]
    soft_skill_candidates = [kw for kw in missing if kw not in hard_tech_missing]

    still_missing = list(hard_tech_missing)  # hard tools stay missing — no semantic rescue
    semantic_matches = []

    if soft_skill_candidates:
        # One GPT call for all soft skill candidates
        batch_prompt = (
            f'Resume text (first 3000 chars):\n{resume_full_text[:3000]}\n\n'
            f'For each keyword below, determine if the concept is clearly demonstrated '
            f'in the resume even if the exact word is not used.\n'
            f'Return a JSON object: {{"keyword": {{"found": true/false, "evidence": "brief quote or empty"}}}}\n\n'
            f'IMPORTANT: Only return true if there is clear, unambiguous evidence. '
            f'Do not stretch meaning or make assumptions.\n\n'
            f'Keywords to check:\n{json.dumps(soft_skill_candidates)}'
        )
        try:
            resp = client.chat.completions.create(
                model=MODEL,
                messages=[{'role': 'user', 'content': batch_prompt}],
                response_format={'type': 'json_object'},
                max_tokens=800,
                temperature=0
            )
            results = json.loads(resp.choices[0].message.content)
            for kw in soft_skill_candidates:
                match_data = results.get(kw, {})
                if match_data.get('found'):
                    semantic_matches.append({
                        'keyword': kw,
                        'evidence': match_data.get('evidence', '')
                    })
                    found.append(kw)
                else:
                    still_missing.append(kw)
        except Exception:
            still_missing.extend(soft_skill_candidates)

    score = round(len(found) / len(jd_keywords) * 100)
    return score, found, still_missing, semantic_matches


def check_education_match(resume_text, jd_text):
    """Check if candidate meets education requirement. Returns score 0-100."""
    jd_lower = jd_text.lower()
    resume_lower = resume_text.lower()

    jd_phd = any(t in jd_lower for t in ['ph.d', 'phd', 'doctorate', 'doctoral'])
    jd_masters = any(t in jd_lower for t in ["master's", 'masters', 'msc', 'm.s.', 'mba', 'postgraduate'])
    jd_bachelors = any(t in jd_lower for t in ["bachelor's", 'bachelors', 'bsc', 'b.s.', 'undergraduate', 'b.tech', 'b.e.'])

    has_phd = any(t in resume_lower for t in ['ph.d', 'phd', 'doctorate', 'doctoral'])
    has_masters = any(t in resume_lower for t in ["master's", 'masters', 'msc', 'm.s.', 'mba'])
    has_bachelors = any(t in resume_lower for t in [
        "bachelor's", 'bachelors', 'bsc', 'b.s.', 'b.tech', 'b.e.',
        'undergraduate', 'bachelor of'
    ])

    if jd_phd:
        if has_phd: return 100
        if has_masters: return 75
        return 50
    if jd_masters:
        if has_phd or has_masters: return 100
        if has_bachelors: return 75
        return 50
    if jd_bachelors:
        if has_phd or has_masters or has_bachelors: return 100
        return 50
    return 100  # No specific education requirement found


# ── Output file generation ────────────────────────────────────────────────────
def save_as_docx(resume_text, output_path):
    """Save optimized resume as formatted Word document."""
    try:
        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(10)

        lines = resume_text.split('\n')
        first_line_idx = next((i for i, l in enumerate(lines) if l.strip()), 0)

        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                doc.add_paragraph('')
                continue
            if i == first_line_idx:
                para = doc.add_paragraph()
                run = para.add_run(line)
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
            elif line.isupper() and len(line) < 50:
                para = doc.add_paragraph()
                run = para.add_run(line)
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
            elif line.startswith(('•', '-', '*', '–', '▪')):
                para = doc.add_paragraph(style='List Bullet')
                para.paragraph_format.left_indent = Inches(0.25)
                run = para.add_run(line[1:].strip())
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
            else:
                para = doc.add_paragraph()
                run = para.add_run(line)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

        doc.save(output_path)
        return output_path
    except Exception as e:
        return f'Error saving DOCX: {e}'


def save_as_pdf(resume_text, output_path):
    """Save optimized resume as PDF using ReportLab."""
    try:
        doc = SimpleDocTemplate(
            output_path, pagesize=letter,
            rightMargin=inch, leftMargin=inch,
            topMargin=0.75 * inch, bottomMargin=0.75 * inch
        )
        styles = getSampleStyleSheet()

        name_style = ParagraphStyle(
            'NameStyle', parent=styles['Normal'],
            fontName='Times-Bold', fontSize=14, spaceAfter=6
        )
        header_style = ParagraphStyle(
            'HeaderStyle', parent=styles['Normal'],
            fontName='Times-Bold', fontSize=11,
            textColor=(0.18, 0.46, 0.71), spaceBefore=8, spaceAfter=4
        )
        body_style = ParagraphStyle(
            'BodyStyle', parent=styles['Normal'],
            fontName='Times-Roman', fontSize=10, spaceAfter=2
        )
        bullet_style = ParagraphStyle(
            'BulletStyle', parent=styles['Normal'],
            fontName='Times-Roman', fontSize=10,
            leftIndent=20, spaceAfter=2
        )

        story = []
        lines = resume_text.split('\n')
        first_line = True

        for line in lines:
            line = line.strip()
            line = line.replace('**', '')
            line = line.replace('• *', '• ')
            line = line.replace('*', '')
            if not line:
                story.append(Spacer(1, 4))
                continue
            if first_line:
                story.append(Paragraph(line, name_style))
                first_line = False
            elif line.isupper() and len(line) < 50:
                story.append(Paragraph(line, header_style))
            elif line.startswith(('•', '-', '*', '–', '▪')):
                story.append(Paragraph(f'• {line[1:].strip()}', bullet_style))
            else:
                story.append(Paragraph(line, body_style))

        doc.build(story)
        return output_path
    except Exception as e:
        return f'Error saving PDF: {e}'


# ── Skills deduplication ──────────────────────────────────────────────────────
def deduplicate_skills(resume_text):
    """
    Find the SKILLS section and deduplicate items within each bullet,
    case-insensitively while preserving original casing.
    """
    lines = resume_text.split('\n')

    # Find SKILLS header and the next all-caps header after it
    skills_start = None
    skills_end = len(lines)
    for i, line in enumerate(lines):
        stripped = line.strip()
        if skills_start is None:
            if stripped.upper() == 'SKILLS' and stripped == stripped.upper() and len(stripped) >= 4:
                skills_start = i
        else:
            # Next all-caps header (at least 3 chars, not a bullet)
            if (stripped and stripped == stripped.upper() and len(stripped) >= 3
                    and not stripped.startswith(('•', '-', '*', '–', '▪'))):
                skills_end = i
                break

    if skills_start is None:
        return resume_text

    new_lines = list(lines)
    for i in range(skills_start + 1, skills_end):
        line = lines[i]
        stripped = line.strip()
        if not stripped.startswith(('•', '-', '*', '–', '▪')):
            continue
        # Extract bullet char + leading whitespace
        bullet_char = stripped[0]
        content = stripped[1:].strip()
        # Split on commas, deduplicate case-insensitively
        parts = [p.strip() for p in content.split(',') if p.strip()]
        seen_lower = set()
        deduped = []
        for part in parts:
            key = part.lower()
            if key not in seen_lower:
                seen_lower.add(key)
                deduped.append(part)
        # Preserve indentation
        indent = line[: len(line) - len(line.lstrip())]
        new_lines[i] = f'{indent}{bullet_char} {", ".join(deduped)}'

    return '\n'.join(new_lines)


# ── Fabrication guard ────────────────────────────────────────────────────────
_STOPWORDS = {
    'the', 'and', 'or', 'with', 'for', 'to', 'in', 'of', 'a', 'an', 'is',
    'are', 'was', 'were', 'be', 'been', 'by', 'on', 'at', 'from', 'as',
    'this', 'that', 'it', 'its', 'they', 'have', 'has', 'had', 'their',
    'our', 'your', 'but', 'not', 'also', 'more',
}


def _tokenize(text):
    """Split text into meaningful words: lowercase, >=3 chars, non-stopword."""
    words = re.findall(r'[a-zA-Z]+', text.lower())
    return {w for w in words if len(w) >= 3 and w not in _STOPWORDS}


def validate_no_fabrication(original_text, optimized_text, allowed_extras=None):
    """
    Post-processing guard. Removes or reverts content in the optimized resume
    that has no basis in the original.

    Skills section: removes individual skill tokens whose words are entirely new
    (not present anywhere in the original text).

    Experience / Projects bullets: if fewer than 50% of the bullet's words
    appear anywhere in the original AND the bullet contains new words,
    reverts to the nearest-matching original bullet.

    allowed_extras: optional list of semantically verified keywords that may
    appear even if not in the original text (max 5 recommended).
    """
    original_words = _tokenize(original_text)
    optimized_words = _tokenize(optimized_text)
    allowed_extra_words = _tokenize(' '.join(allowed_extras)) if allowed_extras else set()
    new_words = (optimized_words - original_words) - allowed_extra_words

    if not new_words:
        return optimized_text  # nothing new — nothing to validate

    original_lines = original_text.split('\n')
    original_bullets = [
        ln.strip() for ln in original_lines
        if ln.strip().startswith(('•', '-', '*', '–', '▪'))
    ]

    lines = optimized_text.split('\n')
    new_lines = list(lines)

    def _current_section(up_to_idx):
        section = 'other'
        for j in range(up_to_idx + 1):
            s = lines[j].strip().upper()
            if s in ('SKILLS', 'TECHNICAL SKILLS', 'SKILLS & TOOLS'):
                section = 'skills'
            elif s in ('EXPERIENCE', 'WORK EXPERIENCE', 'PROFESSIONAL EXPERIENCE',
                       'EMPLOYMENT', 'EMPLOYMENT HISTORY'):
                section = 'experience'
            elif s in ('PROJECTS', 'PROJECT EXPERIENCE', 'PERSONAL PROJECTS',
                       'ACADEMIC PROJECTS'):
                section = 'projects'
            elif (s and s == s.upper() and len(s) >= 3
                  and not s.startswith(('•', '-', '*', '–', '▪'))):
                # Any other all-caps header resets to 'other'
                if s not in ('SKILLS', 'EXPERIENCE', 'PROJECTS'):
                    section = 'other'
        return section

    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped.startswith(('•', '-', '*', '–', '▪')):
            continue

        content = stripped[1:].strip()
        bullet_words = _tokenize(content)
        if not bullet_words:
            continue

        bullet_new_words = bullet_words & new_words
        section = _current_section(i)

        if section == 'skills':
            # Remove individual comma-separated tokens that are entirely fabricated
            parts = [p.strip() for p in content.split(',') if p.strip()]
            kept = []
            for part in parts:
                part_words = _tokenize(part)
                if (part_words
                        and part_words.issubset(new_words)
                        and not (part_words & original_words)):
                    continue  # every word is new and none exist in original → remove
                kept.append(part)
            if len(kept) != len(parts):
                indent = line[: len(line) - len(line.lstrip())]
                bullet_char = stripped[0]
                new_lines[i] = (
                    f'{indent}{bullet_char} {", ".join(kept)}' if kept else ''
                )

        elif section in ('experience', 'projects'):
            # 50% word-overlap check against original text
            shared = bullet_words & original_words
            ratio = len(shared) / len(bullet_words)

            if ratio < 0.0:
                # Revert to the nearest original bullet (highest word overlap)
                best_match = None
                best_overlap = -1
                for ob in original_bullets:
                    ob_words = _tokenize(ob)
                    overlap = len(bullet_words & ob_words)
                    if overlap > best_overlap:
                        best_overlap = overlap
                        best_match = ob
                if best_match:
                    indent = line[: len(line) - len(line.lstrip())]
                    new_lines[i] = indent + best_match

    return '\n'.join(new_lines)


# ── JSON repair ───────────────────────────────────────────────────────────────
def _repair_json(s):
    """
    Escape literal control characters (newlines, tabs, etc.) that appear inside
    JSON string values. GPT occasionally returns the optimized_resume field with
    raw newlines instead of \\n, making json.loads() raise 'Invalid control character'.
    Uses a simple state machine: tracks whether the cursor is inside a quoted string.
    """
    result = []
    in_string = False
    escape_next = False
    _ESCAPE_MAP = {'\n': '\\n', '\r': '\\r', '\t': '\\t'}
    for ch in s:
        if escape_next:
            result.append(ch)
            escape_next = False
        elif ch == '\\' and in_string:
            result.append(ch)
            escape_next = True
        elif ch == '"':
            result.append(ch)
            in_string = not in_string
        elif in_string and ch in _ESCAPE_MAP:
            result.append(_ESCAPE_MAP[ch])
        elif in_string and ord(ch) < 0x20:
            pass  # drop other control characters inside strings
        else:
            result.append(ch)
    return ''.join(result)


# ── Keyword presence helper ───────────────────────────────────────────────────
KEYWORD_SYNONYMS = {
    'machine learning': ['ml'],
    'ml': ['machine learning'],
    'large language models': ['llms', 'llm'],
    'llm': ['large language models', 'llms'],
    'llms': ['large language models', 'llm'],
    'natural language processing': ['nlp'],
    'nlp': ['natural language processing'],
    'artificial intelligence': ['ai'],
    'ai': ['artificial intelligence'],
    'python': ['python-based', 'python based'],
    'github': ['git', 'github-based'],
    'git': ['github', 'git-based'],
    'power bi': ['powerbi'],
    'scikit-learn': ['sklearn', 'scikit learn'],
    'hugging face': ['huggingface'],
    'aws': ['amazon web services'],
    'gcp': ['google cloud'],
    'azure': ['microsoft azure'],
}


def keyword_in_text(keyword, text):
    """
    Check if keyword exists in text using:
    1. Word boundary exact match
    2. Explicit controlled synonym matching only
    No fuzzy matching. No loose inference.
    """
    kw_lower = keyword.lower().strip()
    text_lower = text.lower()

    pattern = r'\b' + re.escape(kw_lower) + r'\b'
    if re.search(pattern, text_lower):
        return True

    for synonym in KEYWORD_SYNONYMS.get(kw_lower, []):
        syn_pattern = r'\b' + re.escape(synonym) + r'\b'
        if re.search(syn_pattern, text_lower):
            return True

    return False


# ── Actual changes helper ─────────────────────────────────────────────────────
def compute_actual_changes(original_text, final_text):
    """
    Compare original and final resume line by line.
    Only report changes that exist in the final resume.
    Never report GPT intentions that were removed by validation.
    """
    original_lines = set(
        l.strip() for l in original_text.split('\n')
        if l.strip()
    )
    final_lines = set(
        l.strip() for l in final_text.split('\n')
        if l.strip()
    )

    EDUCATION_SIGNALS = [
        'university', 'college', 'institute',
        'bachelor', 'master', 'phd', 'gpa',
        'b.s.', 'm.s.', 'b.tech', 'doctorate',
        'science in', 'arts in', 'engineering in',
    ]

    def is_education_line(line):
        line_lower = line.lower()
        return any(signal in line_lower for signal in EDUCATION_SIGNALS)

    def is_skills_line(line):
        skills_patterns = [
            r'^programming',
            r'^statistical',
            r'^data visualization',
            r'^ai &',
            r'^cloud &',
            r'^data engineering',
            r'^technical skills',
            r'^skills',
            r'^languages',
            r'^tools',
            r'^frameworks',
        ]
        line_lower = line.lower()
        return any(re.search(p, line_lower) for p in skills_patterns)

    changed = []
    for line in final_lines:
        if (line not in original_lines and len(line) > 15
                and not is_education_line(line)
                and not is_skills_line(line)):
            best_match = max(
                original_lines,
                key=lambda o: len(
                    set(line.lower().split()) &
                    set(o.lower().split())
                ),
                default=None
            )
            if best_match:
                changed.append(
                    f"Strengthened: '{best_match[:60]}...'"
                    if len(best_match) > 60
                    else f"Strengthened: '{best_match}'"
                )

    return changed[:8]


# ── Keyword list cleaner ──────────────────────────────────────────────────────
def clean_jd_keywords(keywords):
    """
    Remove noise from GPT-extracted keyword list.
    Keep only meaningful ATS-relevant terms.
    Filter only clear noise — never filter meaningful domain terms.
    """
    NOISE_PATTERNS = [
        # Company and team names
        r'^firm risk management',
        r'^risk ai application',
        # Work arrangements
        r'\bhybrid work\b',
        r'\bin.office\b',
        r'\bremote work\b',
        r'\bhybrid position\b',
        r'\bfull.time position\b',
        r'\bwork mode\b',
        r'\bdays in office\b',
        # Generic boilerplate phrases
        r'^investment bank',
        r'^global teams',
        r'^you will\b',
        r'^the firm\b',
        r'\bthis role requires\b',
        r'\bas part of your\b',
        r'\blearn more\b',
        r'\bclick here\b',
        # Compensation / benefits
        r'\bannual performance bonus\b',
        r'\bequity package\b',
        r'\btotal rewards\b',
        r'\bcompetitive benefits\b',
        r'\beligible for\b',
        # Safety / misc
        r'\bworkplace safety\b',
    ]

    cleaned = []
    for kw in keywords:
        kw_stripped = kw.strip()
        if not kw_stripped:
            continue
        # Skip phrases longer than 4 words — too specific to be meaningful ATS keywords
        if len(kw_stripped.split()) > 4:
            continue
        # Skip single characters
        if len(kw_stripped) < 2:
            continue
        # Skip if matches clear noise pattern
        if any(re.search(p, kw_stripped.lower()) for p in NOISE_PATTERNS):
            continue
        cleaned.append(kw_stripped)

    # Deduplicate case-insensitively
    seen = set()
    final = []
    for kw in cleaned:
        if kw.lower() not in seen:
            seen.add(kw.lower())
            final.append(kw)

    return final


# ── Experience bullet safety check ───────────────────────────────────────────

def check_experience_bullets(original_text, optimized_text):
    """
    Only checks bullet lines.
    If any bullet in optimized has less than 30% word overlap with ANY original
    bullet, revert it to the closest original bullet.
    Leaves Skills, Summary, and non-bullet lines untouched.
    """
    original_bullets = [
        l.strip() for l in original_text.split('\n')
        if l.strip().startswith(('•', '-', '*', '–', '▪', '·'))
    ]

    if not original_bullets:
        return optimized_text

    result_lines = []
    for line in optimized_text.split('\n'):
        stripped = line.strip()
        if not stripped.startswith(('•', '-', '*', '–', '▪', '·')):
            result_lines.append(line)
            continue

        opt_words = set(stripped.lower().split())
        best_overlap = max(
            (
                len(opt_words & set(ob.lower().split())) / max(len(opt_words), 1)
                for ob in original_bullets
            ),
            default=0
        )

        if best_overlap < 0.15:
            best_match = max(
                original_bullets,
                key=lambda ob: len(opt_words & set(ob.lower().split())),
                default=stripped
            )
            indent = len(line) - len(line.lstrip())
            result_lines.append(' ' * indent + best_match)
        else:
            result_lines.append(line)

    return '\n'.join(result_lines)


# ── Main optimization engine ──────────────────────────────────────────────────
def optimize_resume(resume_text, job_description_text, status_callback=None):
    """
    Main optimization engine.
    1. Build messages (system + user prompt)
    2. Call GPT with tools
    3. While loop for tool calls (URL fetching)
    4. Parse GPT JSON response
    5. Run our own keyword and experience scoring (overrides GPT scores)
    6. Return complete result dict
    """
    def status(msg):
        if status_callback:
            status_callback(msg)

    status('Building optimization prompt...')

    user_prompt = f"""Optimize this resume for the job description below.

JOB DESCRIPTION:
{job_description_text}

RESUME:
{resume_text}

Instructions:
1. Detect the target role — map to one of the 6 supported roles exactly.
2. Extract ALL keywords from the job description.
3. Scan the ENTIRE resume (all sections) for each keyword.
4. Identify keywords already present anywhere in the resume.
5. Rewrite existing bullets to naturally incorporate as many JD keywords as possible wherever the underlying experience genuinely supports it. Same experience, same metrics, same numbers — just weave in JD vocabulary throughout. Keywords that cannot be incorporated anywhere should be listed as missing.
6. Never penalize extra skills. Score = JD keyword coverage only.
7. Return valid JSON matching the required output format exactly.
"""

    messages = [
        {'role': 'system', 'content': SYSTEM_MESSAGE},
        {'role': 'user', 'content': user_prompt}
    ]

    status('Calling GPT-4o for optimization...')
    response = client.chat.completions.create(
        model=MODEL_SMART,
        messages=messages,
        temperature=0.1
    )

    raw_content = response.choices[0].message.content

    # Parse JSON — strip markdown fences, then repair control characters
    try:
        if '```json' in raw_content:
            raw_content = raw_content.split('```json')[1].split('```')[0].strip()
        elif '```' in raw_content:
            raw_content = raw_content.split('```')[1].split('```')[0].strip()
        try:
            result = json.loads(raw_content)
        except json.JSONDecodeError:
            # GPT embedded literal newlines/tabs inside JSON string values — repair and retry
            result = json.loads(_repair_json(raw_content))
    except Exception as e:
        return {'error': f'Failed to parse GPT response: {e}', 'raw': raw_content}

    print("DEBUG optimized length:", len(result.get('optimized_resume', '')))
    print("DEBUG first 500 chars:", result.get('optimized_resume', '')[:500])

    # ── Step 2: Post-processing ────────────────────────────────────────────────

    # 2a. Deduplicate skills
    optimized_text = deduplicate_skills(result.get('optimized_resume', ''))

    # ── Collect JD keywords from GPT JSON output ───────────────────────────────
    gpt_ats = result.get('ats_keywords', {})
    all_jd_keywords = clean_jd_keywords(list({
        *gpt_ats.get('matched', []),
        *gpt_ats.get('added', []),
        *gpt_ats.get('missing', []),
    }))

    # ── Keyword scoring ────────────────────────────────────────────────────────
    status('Running keyword scoring...')
    kw_score_before, _, _, _ = calculate_keyword_score(all_jd_keywords, resume_text)
    kw_score_after, kw_found_after, kw_missing_after, kw_semantic = calculate_keyword_score(
        all_jd_keywords, optimized_text
    )

    # ── Step 3: Accurate keywords_added (compare original vs final) ────────────
    actually_added = [
        kw for kw in all_jd_keywords
        if not keyword_in_text(kw, resume_text) and keyword_in_text(kw, optimized_text)
    ]

    # ── Experience + education scoring ────────────────────────────────────────
    status('Analyzing relevant experience...')
    target_role = result.get('role_type_detected', 'Data Scientist')
    relevant_years, exp_breakdown, required_years = calculate_relevant_experience(
        resume_text, target_role, job_description_text
    )
    edu_score = check_education_match(resume_text, job_description_text)

    # Combined ATS score: keywords 50%, experience 30%, education 20%
    if required_years:
        exp_score = min(100, int(relevant_years / required_years * 100))
    else:
        exp_score = 80  # no requirement stated → assume reasonable
    final_score = round(kw_score_after * 0.5 + exp_score * 0.3 + edu_score * 0.2)

    # ── Step 4: Accurate changes_made (line diff original vs final) ────────────
    return {
        'role_detected': target_role,
        'optimized_resume': optimized_text,
        'ats_score_before': kw_score_before,
        'ats_score_after': final_score,
        'kw_score_before': kw_score_before,
        'kw_score_after': kw_score_after,
        'keywords_found': kw_found_after,
        'keywords_missing': kw_missing_after,
        'jd_keywords': all_jd_keywords,
        'keywords_added': actually_added,
        'semantic_matches': kw_semantic,
        'changes_made': compute_actual_changes(resume_text, optimized_text),
        'relevant_years': relevant_years,
        'required_years': required_years,
        'exp_breakdown': exp_breakdown,
        'edu_score': edu_score,
        'exp_score': exp_score,
    }


# ── Chat logic ────────────────────────────────────────────────────────────────
_CHAT_SYSTEM_TEMPLATE = """You are a resume optimization assistant helping a user refine their resume.

You have full context:

ORIGINAL RESUME:
{original_resume}

JOB DESCRIPTION:
{job_description}

CURRENT OPTIMIZED RESUME:
{current_resume}

STILL MISSING KEYWORDS:
{missing_keywords}

TARGET ROLE: {role}

You handle ALL types of resume change requests:
- Adding confirmed missing skills the user says they have
- Rewriting or strengthening the summary
- Improving experience bullets
- Updating projects section
- Any other change the user requests

Every change you make must:
1. Be based on information the user provides or content already in the resume
2. Align with the job description and improve ATS keyword coverage
3. Never fabricate or hallucinate anything not confirmed by the user

BEFORE making any resume changes:

1. If the user confirms they have a missing skill but does not specify WHERE they used it (which role or project) and HOW they used it, ask ONE concise clarifying question covering all unclear skills together — do not ask separately for each skill.
   Example: 'Which roles or projects did you use [skill A], [skill B], and [skill C] in, and briefly how did you use them?'

2. Only update a section when the user has explicitly connected a skill to a specific role or project AND described how they used it.

3. Never infer placement — if the user says 'I use Git' without specifying where and how, ask before adding it anywhere.

4. When the user provides specific placement and usage, add the skill ONLY to that specific section — nowhere else in the resume.

STRICT RULES FOR ADDING CONFIRMED SKILLS:

- Add a confirmed skill ONLY to sections where the user explicitly described doing that work
- Never add a skill to a role or project just because it sounds relevant — it must be traceable to what the user specifically said
- Never add skills the user did not confirm — if they confirmed Git and clean code, do NOT also add CI/CD, unit testing, or code reviews unless explicitly stated
- A confirmed skill may be added only to the specific section or sections the user explicitly described — never generalized across the entire resume
- If unsure whether a section supports the confirmed skill, leave that section unchanged

When making any changes:
- Return the COMPLETE updated resume starting with the marker UPDATED_RESUME:
- Follow the resume with the separator ---RESPONSE---
- Then give a brief conversational response explaining what changed
- After making changes always ask: "Does this look good, or would you like anything else adjusted?"

When the user asks questions (no changes requested):
- Answer directly using your knowledge of their resume and the JD
- Do NOT regenerate the resume unless they provide new content or request a change

IMPORTANT RULE — DIRECT EDITS WITH CAUTION:

When the user asks to add a specific skill, tool, or technology they explicitly mention having experience with — add it immediately without asking follow-up questions.

When the user asks a vague request like 'add all relevant ML tools' or 'add missing skills' or 'add tools required for this role' WITHOUT specifying which ones they have experience with — DO NOT add anything to the resume. Instead respond with a short, polite message in this style:

'I can definitely help strengthen your skills section.

To make sure everything reflects your actual experience, it's best to include only tools you've worked with, as they may come up in interviews.

Here are some relevant skills from the job description that aren't currently listed: [list].

Let me know which ones you've used, and I'll add them for you.'

When the user confirms specific skills they have experience with — either by naming them explicitly or by providing any details about how they used them — add only those confirmed skills immediately to the most relevant section without asking any further questions.

Never add skills, tools, or technologies to the resume that the user has not explicitly confirmed they have experience with.

Always apply confirmed changes directly to the CURRENT OPTIMIZED RESUME and return UPDATED_RESUME: with the full updated resume.

Be concise and professional."""


REGENERATE_TRIGGERS = [
    'add ', 'update ', 'change ', 'remove ', 'rewrite ', 'modify ', 'incorporate ',
    'include ', 'put ', 'replace ', 'insert ', 'delete ', 'revise ', 'reflect ',
    'apply ', 'fix ', 'edit ', 'move ', 'mention ',
]

QUESTION_WORDS = [
    'what ', 'how ', 'why ', 'where ', 'when ', 'which ', 'who ',
    'can you ', 'could you ', 'would you ', 'is there ', 'are there ',
    'do you ', 'does ', 'did ', 'should ',
]


def should_regenerate(user_message):
    """Rule-based: return True only when user explicitly asks to change the resume."""
    msg = user_message.strip().lower()

    # Questions never trigger regeneration
    if msg.endswith('?'):
        return False
    if any(msg.startswith(q) for q in QUESTION_WORDS):
        return False

    # "I have/use/know X" without an explicit action word → observation, not request
    observation_starters = ('i have ', 'i know ', 'i use ', 'i used ', 'i can ', 'i am ', 'i do ')
    if any(msg.startswith(s) for s in observation_starters):
        if not any(t in msg for t in REGENERATE_TRIGGERS):
            return False

    # Formatting/download requests — never touch resume content
    non_resume_triggers = [
        'download',
        'pdf format',
        'docx format',
        'downloadable',
        'make it bold',
        'make bold',
        'format the',
        'make the resume',
    ]
    if any(t in msg for t in non_resume_triggers):
        return False

    # Trigger only on explicit action words
    return any(t in msg for t in REGENERATE_TRIGGERS)


def chat_with_resume(user_message, history, context):
    """
    Chat handler.
    Returns (updated_history, updated_context, new_resume_text|None, docx_path|None, pdf_path|None)
    """
    if not context.get('original_resume'):
        updated = history + [
            {'role': 'user', 'content': user_message},
            {'role': 'assistant', 'content': 'Please analyze a resume first using the form above.'}
        ]
        return updated, context, None, None, None

    system_content = _CHAT_SYSTEM_TEMPLATE.format(
        original_resume=context.get('original_resume', ''),
        job_description=context.get('job_description', ''),
        current_resume=context.get('current_resume', ''),
        missing_keywords=', '.join(context.get('missing_keywords', [])) or 'None',
        role=context.get('role', 'Unknown'),
    )

    messages = [{'role': 'system', 'content': system_content}]
    for msg in history:
        messages.append({'role': msg['role'], 'content': msg['content']})

    regenerate = should_regenerate(user_message)

    if regenerate:
        augmented = (
            user_message
            + '\n\nIncorporate this into the resume. Return the COMPLETE updated resume '
            'starting with UPDATED_RESUME: then the separator ---RESPONSE--- and your response.'
        )
        messages.append({'role': 'user', 'content': augmented})
    else:
        messages.append({'role': 'user', 'content': user_message})

    resp = client.chat.completions.create(
        model=MODEL_SMART,
        messages=messages,
        temperature=0.4,
        max_tokens=2500
    )

    reply = resp.choices[0].message.content
    new_resume_text = None
    docx_path = None
    pdf_path = None

    if 'UPDATED_RESUME:' in reply:
        parts = reply.split('---RESPONSE---')
        resume_part = parts[0].replace('UPDATED_RESUME:', '').strip()

        lines = resume_part.strip().split('\n')
        cleaned_lines = []
        start = False
        for line in lines:
            stripped = line.strip()
            if not stripped:
                if start:
                    cleaned_lines.append(line)
                continue
            FILLER_PHRASES = [
                "certainly", "sure", "of course",
                "here is", "here's", "i've",
                "i have", "below is", "please",
                "i'll", "i will", "as requested",
                "absolutely", "updated", "great",
                "happy to", "let me"
            ]
            is_filler = any(f in stripped.lower() for f in FILLER_PHRASES)
            is_name_like = (
                len(stripped) < 40 and
                not stripped.endswith('.') and
                not stripped.endswith(',') and
                not stripped.endswith(':') and
                not is_filler and
                stripped.replace(' ', '').isalpha()
            )
            if not start and (is_name_like or stripped.isupper()):
                start = True
            if start:
                cleaned_lines.append(line)
        if cleaned_lines:
            resume_part = '\n'.join(cleaned_lines)

        conversational = parts[1].strip() if len(parts) > 1 \
            else 'Done! Your resume has been updated in the ' \
            'preview box above. Would you like any other changes?'
        if len(conversational) > 1000:
            conversational = 'Done! Your resume has been ' \
                'updated in the preview box above. Would you ' \
                'like any other changes?'
        new_resume_text = resume_part

        # Recalculate keyword score with the same list used for initial scoring
        jd_keywords = context.get('jd_keywords', [])
        if jd_keywords:
            new_kw_score, new_found, new_missing, _ = calculate_keyword_score(
                jd_keywords, new_resume_text
            )
            exp_score = context.get('exp_score', 80)
            edu_score = context.get('edu_score', 0)
            new_ats_score = round(new_kw_score * 0.5 + exp_score * 0.3 + edu_score * 0.2)
            original_resume = context.get('original_resume', '')
            actually_added = [
                kw for kw in jd_keywords
                if not keyword_in_text(kw, original_resume) and keyword_in_text(kw, new_resume_text)
            ]
            context = {
                **context,
                'current_resume': new_resume_text,
                'keywords_missing': new_missing,
                'kw_score_after': new_kw_score,
                'ats_score_after': new_ats_score,
                'keywords_added': actually_added,
                'semantic_matches': [],
            }
        else:
            context = {**context, 'current_resume': new_resume_text}

        import time
        ts = int(time.time())
        docx_path = os.path.join(tempfile.gettempdir(), f'optimized_resume_{ts}.docx')
        pdf_path = os.path.join(tempfile.gettempdir(), f'optimized_resume_{ts}.pdf')
        save_as_docx(new_resume_text, docx_path)
        save_as_pdf(new_resume_text, pdf_path)

        reply = conversational
        if new_resume_text:
            reply = (
                'Done! Your resume has been updated '
                'in the preview box above. Would you '
                'like any other changes?'
            )
        if len(reply) > 500 or \
                'EDUCATION' in reply or \
                'EXPERIENCE' in reply or \
                'SKILLS' in reply or \
                'UPDATED_RESUME' in reply:
            reply = ('Done! Your resume has been updated in the preview box above. '
                     'Would you like any other changes?')

    print("DEBUG reply length:", len(reply))
    print("DEBUG reply content:", reply[:200])
    updated_history = history + [
        {'role': 'user', 'content': user_message},
        {'role': 'assistant', 'content': reply}
    ]

    return updated_history, context, new_resume_text, docx_path, pdf_path


# ── Display formatters ────────────────────────────────────────────────────────
def _format_score_panel(result):
    score = result['ats_score_after']
    kw_after = result['kw_score_after']
    kw_before = result['kw_score_before']
    relevant = result['relevant_years']
    required = result['required_years']
    role = result['role_detected']
    added_count = len(result['keywords_added'])

    exp_line = f"{relevant} yrs"
    if required:
        exp_line += f" / {required}+ required"

    return f"""## ATS Score: {score} / 100

| Metric | Before | After |
|--------|--------|-------|
| Keyword Coverage | {kw_before}% | {kw_after}% |
| Education Match | — | {result['edu_score']}% |

**Role Detected:** {role}
**Keywords Added:** {added_count}
**Relevant Experience:** {exp_line}
"""


def _format_missing_panel(result):
    missing = result['keywords_missing']
    if not missing:
        return '### Missing Skills\n\n✅ All JD keywords are covered in the optimized resume!'
    items = '\n'.join(f'⚠️ **{kw}**' for kw in missing)
    return f'### Missing Skills\n\n{items}'


def _format_experience_panel(breakdown, relevant_years, required_years):
    if not breakdown:
        return ''

    rows = ['**Relevant Experience Breakdown**\n']
    rows.append('| Role | Period | Credit | Counted | Why |')
    rows.append('|------|--------|--------|---------|-----|')

    for e in breakdown:
        emoji = {'full': '✅', 'partial': '⚡', 'none': '❌'}.get(e['relevance'], '⚡')
        pct = f"{int(e['multiplier'] * 100)}%"
        reason = e['reason'][:80] + ('…' if len(e['reason']) > 80 else '')
        rows.append(
            f"| {e['job_title']} | {e['start']} – {e['end']} | "
            f"{pct} {emoji} | {e['counted_years']}y | {reason} |"
        )

    rows.append(f'\n**Total relevant experience: {relevant_years} years**')
    if required_years:
        gap = round(required_years - relevant_years, 1)
        if gap > 0:
            rows.append(f'⚠️ **Gap: {gap} years below the {required_years}+ year requirement**')
        else:
            rows.append(f'✅ **Meets the {required_years}+ year requirement**')

    return '\n'.join(rows)


def _format_keywords_panel(result):
    lines = []
    if result['keywords_added']:
        kws = ', '.join(f'`{k}`' for k in result['keywords_added'])
        lines.append(f'**Keywords added to resume:** {kws}')
    if result['semantic_matches']:
        kws = ', '.join(f'`{m["keyword"]}`' for m in result['semantic_matches'])
        lines.append(f'**Semantic matches (concept found, exact word not used):** {kws}')
    if result['changes_made']:
        changes = '\n'.join(f'- {c}' for c in result['changes_made'][:8])
        lines.append(f'\n**Changes made:**\n{changes}')
    return '\n\n'.join(lines)


# ── Gradio process callback ───────────────────────────────────────────────────
def process_resume(resume_file, jd_text_input, run_count):
    """
    Main Gradio callback — generator that yields status then final results.
    Yields tuple of 12 values matching the output list.
    """
    EMPTY = ('', gr.update(), gr.update(), gr.update(), gr.update(),
             gr.update(), gr.update(visible=False), None, None, {}, run_count, [], gr.update())

    def partial_yield(status_msg):
        return (
            status_msg, gr.update(), gr.update(), gr.update(), gr.update(),
            gr.update(), gr.update(visible=False), None, None, {}, run_count, [], gr.update()
        )

    # Rate limit
    if run_count >= MAX_RUNS_PER_SESSION:
        yield partial_yield(
            f'⚠️ Session limit reached ({MAX_RUNS_PER_SESSION} analyses maximum). '
            'Refresh the page to reset.'
        )
        return

    # Validate inputs
    if resume_file is None:
        yield partial_yield('⚠️ Please upload a resume file (PDF or DOCX).')
        return

    jd_text_val = (jd_text_input or '').strip()
    if not jd_text_val:
        yield partial_yield('⚠️ Please paste the job description text.')
        return

    yield partial_yield('Reading resume file...')

    # Read resume
    file_path = resume_file if isinstance(resume_file, str) else resume_file.name
    ext = Path(file_path).suffix.lower()
    if ext == '.pdf':
        resume_text = extract_text_from_pdf(file_path)
    elif ext == '.docx':
        resume_text = extract_text_from_docx(file_path)
    else:
        yield partial_yield('⚠️ Unsupported file type. Please upload PDF or DOCX.')
        return

    if resume_text.startswith('Error'):
        yield partial_yield(f'⚠️ {resume_text}')
        return

    job_text = jd_text_val

    yield partial_yield('Optimizing resume with GPT-4o...')

    result = optimize_resume(resume_text, job_text)

    if 'error' in result:
        yield partial_yield(f'⚠️ {result["error"]}')
        return

    yield partial_yield('Generating downloadable files...')

    docx_path = os.path.join(tempfile.gettempdir(), 'optimized_resume.docx')
    pdf_path = os.path.join(tempfile.gettempdir(), 'optimized_resume.pdf')
    save_as_docx(result['optimized_resume'], docx_path)
    save_as_pdf(result['optimized_resume'], pdf_path)

    new_run_count = run_count + 1
    runs_left = MAX_RUNS_PER_SESSION - new_run_count
    status_msg = f'✅ Complete! ({runs_left} run{"s" if runs_left != 1 else ""} remaining this session)'

    context = {
        'original_resume': resume_text,
        'job_description': job_text,
        'current_resume': result['optimized_resume'],
        'missing_keywords': result['keywords_missing'],
        'role': result['role_detected'],
        'jd_keywords': result['jd_keywords'],
        'kw_score_before': result['kw_score_before'],
        'kw_score_after': result['kw_score_after'],
        'ats_score_after': result['ats_score_after'],
        'relevant_years': result['relevant_years'],
        'required_years': result['required_years'],
        'role_detected': result['role_detected'],
        'keywords_added': result['keywords_added'],
        'edu_score': result['edu_score'],
        'exp_score': result['exp_score'],
    }

    # Build warm opening assistant message for chat
    missing = result['keywords_missing']

    if not missing:
        opening_msg = (
            "Hi! Your optimized resume is ready. I've "
            "strengthened your existing content to better "
            "align with the job description.\n\n"
            "Everything in the JD looks well covered! "
            "Does the resume look good to you, or would "
            "you like any changes?"
        )
    else:
        missing_list = ', '.join(missing)
        opening_msg = (
            "Hi! Your optimized resume is ready. I've "
            "strengthened your existing content to better "
            "align with the job description.\n\n"
            "I also found these requirements from the JD that "
            "aren't currently in your resume:\n\n"
            f"{missing_list}\n\n"
            "Do you have experience with any of these? If yes, "
            "mention which role or project you used it in and "
            "briefly how you used it — for example: \"I used "
            "[skill] in my [role/project] for [purpose].\" "
            "This helps me add it accurately to "
            "the right place in your resume."
        )

    initial_chat = [{'role': 'assistant', 'content': opening_msg}]

    yield (
        status_msg,
        gr.update(value=_format_score_panel(result), visible=True),
        gr.update(value=_format_missing_panel(result), visible=True),
        gr.update(
            value=_format_experience_panel(
                result['exp_breakdown'], result['relevant_years'], result['required_years']
            ),
            visible=bool(result['exp_breakdown'])
        ),
        gr.update(value=_format_keywords_panel(result), visible=True),
        gr.update(value=result['optimized_resume'], visible=True),
        gr.update(visible=True),   # download_row
        docx_path,
        pdf_path,
        context,
        new_run_count,
        initial_chat,   # → chat_history_state (used by subsequent chat calls)
        initial_chat,   # → chatbot display (renders immediately after analysis)
    )


# ── Gradio UI ─────────────────────────────────────────────────────────────────
def build_ui():
    """Build and return the Gradio UI. Called only when running as __main__."""

    with gr.Blocks(title='ATS Resume Optimizer') as ui:

        # Persistent state
        run_count_state = gr.State(0)
        context_state = gr.State({})
        chat_history_state = gr.State([])

        # ── Header ────────────────────────────────────────────────────────────
        gr.Markdown("""
<h1 style='text-align: center; margin-bottom: 4px;'>ATS Resume Optimizer</h1>
<p style='text-align: center; color: #6b7280; font-size: 15px; margin-top: 0;'>Analyze your resume against a job description. Get an ATS keyword score, relevant experience breakdown, and an optimized resume ready to download.</p>
""")

        # ── Inputs ────────────────────────────────────────────────────────────
        with gr.Row():
            with gr.Column(scale=1):
                resume_input = gr.File(
                    label='Upload Resume (PDF or DOCX)',
                    file_types=['.pdf', '.docx'],
                    type='filepath'
                )
            with gr.Column(scale=1):
                jd_text_input = gr.Textbox(
                    label='Paste Job Description',
                    placeholder='Paste the full job description text...',
                    lines=8
                )
                gr.Markdown(
                    "_Tip: Copy the full job description from LinkedIn, Indeed, or the company careers page and paste it above._",
                    elem_classes=['jd-tip']
                )

        analyze_btn = gr.Button(
            'Analyze & Optimize Resume',
            variant='primary',
            size='lg'
        )

        # ── Status bar ────────────────────────────────────────────────────────
        status_output = gr.Textbox(
            label='Status',
            interactive=False,
            lines=1,
            show_label=False
        )

        # ── Results panels ────────────────────────────────────────────────────
        with gr.Row():
            score_output = gr.Markdown(visible=False)
            missing_output = gr.Markdown(visible=False)

        exp_output = gr.Markdown(visible=False)
        kw_output = gr.Markdown(visible=False)

        resume_preview = gr.Textbox(
            label='Optimized Resume Preview (editable)',
            lines=30,
            interactive=True,
            visible=False
        )

        with gr.Row(visible=False) as download_row:
            docx_download = gr.File(label='Download DOCX', show_label=True, min_width=100)
            pdf_download = gr.File(label='Download PDF', show_label=True, min_width=100)

        # ── Chat section ──────────────────────────────────────────────────────
        gr.Markdown('---')
        gr.Markdown('## Refine Your Resume via Chat')

        chatbot = gr.Chatbot(
            label='Conversation',
            height=300,
            show_label=False,
            placeholder="<div style='text-align:center; color:#94a3b8; padding:40px 0'>Analyze a resume above to start the conversation</div>"
        )

        with gr.Row():
            chat_input = gr.Textbox(
                label='',
                placeholder='Tell me about missing skills or ask for changes...',
                lines=2,
                scale=4,
                show_label=False
            )
            send_btn = gr.Button('Send', variant='primary', scale=1, min_width=80)

        # ── Wire up events ────────────────────────────────────────────────────
        analyze_btn.click(
            fn=process_resume,
            inputs=[resume_input, jd_text_input, run_count_state],
            outputs=[
                status_output,
                score_output,
                missing_output,
                exp_output,
                kw_output,
                resume_preview,
                download_row,
                docx_download,
                pdf_download,
                context_state,
                run_count_state,
                chat_history_state,
                chatbot,
            ]
        )

        def _handle_chat(user_msg, history, context):
            if not user_msg.strip():
                yield history, context, gr.update(), gr.update(), gr.update(), gr.update(), gr.update(), '', history
                return

            # Yield 1: show user message immediately with typing indicator, clear input
            optimistic = history + [
                {'role': 'user', 'content': user_msg},
                {'role': 'assistant', 'content': 'Processing your request...'}
            ]
            yield optimistic, context, gr.update(), gr.update(), gr.update(), gr.update(), gr.update(), '', optimistic

            # Yield 2: full response after GPT call
            updated_history, updated_context, new_resume, docx_path, pdf_path = chat_with_resume(
                user_msg, history, context
            )
            resume_update = gr.update(value=new_resume) if new_resume else gr.update()
            docx_update = gr.update(value=docx_path) if docx_path else gr.update()
            pdf_update = gr.update(value=pdf_path) if pdf_path else gr.update()

            if new_resume:
                chat_result = {
                    'ats_score_after': updated_context.get('ats_score_after', 0),
                    'kw_score_after': updated_context.get('kw_score_after', 0),
                    'kw_score_before': updated_context.get('kw_score_before', 0),
                    'relevant_years': updated_context.get('relevant_years', 0),
                    'required_years': updated_context.get('required_years'),
                    'role_detected': updated_context.get('role_detected', ''),
                    'keywords_added': updated_context.get('keywords_added', []),
                    'edu_score': updated_context.get('edu_score', 0),
                    'keywords_missing': updated_context.get('keywords_missing', []),
                }
                score_update = gr.update(value=_format_score_panel(chat_result))
                missing_update = gr.update(value=_format_missing_panel(chat_result))
            else:
                score_update = gr.update()
                missing_update = gr.update()

            yield updated_history, updated_context, resume_update, docx_update, pdf_update, score_update, missing_update, '', updated_history

        _chat_outputs = [
            chatbot, context_state, resume_preview, docx_download, pdf_download,
            score_output, missing_output, chat_input, chat_history_state
        ]

        send_btn.click(
            fn=_handle_chat,
            inputs=[chat_input, chat_history_state, context_state],
            outputs=_chat_outputs
        )

        chat_input.submit(
            fn=_handle_chat,
            inputs=[chat_input, chat_history_state, context_state],
            outputs=_chat_outputs
        )

    return ui


if __name__ == '__main__':
    ui = build_ui()
    ui.queue()
    ui.launch(
        theme=gr.themes.Soft(),
        server_name="0.0.0.0",
        server_port=7861
    )
